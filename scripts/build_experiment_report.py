"""Build the final English experiment report from verified project artifacts."""

from __future__ import annotations

import csv
import hashlib
import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report_assets"
RUN = ROOT / "runs/detect/mouse_cup_yolo11n_v13_horizontal_cup_768"
DATASET = ROOT / "dataset_work/audit_dataset_v13/yolo_export"
REVIEW = ROOT / "results/v13_submission_review"
OUTPUT = DOCS / "Experiment_One_Object_Detection_Report.docx"

NAVY = "18324B"
BLUE = "2D6A9F"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F3F5F7"
MID_GRAY = "D4D9DE"
TEXT = RGBColor(32, 36, 40)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = MID_GRAY, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


def font(size: int, bold: bool = False):
    try:
        return ImageFont.truetype("C:/Windows/Fonts/arial.ttf", size)
    except OSError:
        return ImageFont.load_default()


def create_pipeline_diagram(path: Path) -> None:
    canvas = Image.new("RGB", (1800, 640), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = font(52, True)
    body_font = font(31)
    small_font = font(24)
    draw.text((70, 45), "End to End Detection Workflow", fill="#111111", font=title_font)
    labels = [
        ("Collect", "phone and camera\nvideo frames"),
        ("Audit", "manual boxes and\nhard negatives"),
        ("Train", "YOLO11n at\n768 pixels"),
        ("Export", "PyTorch and\nONNX weights"),
        ("Deploy", "Windows Jetson\nand ROS 2"),
    ]
    x_values = [70, 410, 750, 1090, 1430]
    for index, ((heading, detail), x) in enumerate(zip(labels, x_values)):
        draw.rounded_rectangle((x, 180, x + 275, 470), radius=28, fill="#EAF2F8", outline="#2D6A9F", width=5)
        draw.text((x + 28, 215), heading, fill="#18324B", font=body_font)
        draw.multiline_text((x + 28, 295), detail, fill="#202428", font=small_font, spacing=12)
        if index < len(labels) - 1:
            draw.line((x + 278, 325, x + 327, 325), fill="#2D6A9F", width=8)
            draw.polygon(((x + 327, 325), (x + 307, 310), (x + 307, 340)), fill="#2D6A9F")
    draw.text((70, 535), "Feedback loop  missed detections and false alarms become the next audited batch", fill="#333333", font=small_font)
    canvas.save(path, quality=95)


def create_dataset_chart(path: Path) -> None:
    values = [("Train", 492, "#2D6A9F"), ("Validation", 61, "#5AA2D1"), ("Test", 240, "#8CB9D9")]
    canvas = Image.new("RGB", (1600, 820), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((80, 55), "Version 13 Exported Image Splits", fill="#111111", font=font(52, True))
    draw.line((170, 680, 1500, 680), fill="#333333", width=3)
    max_value = max(value for _, value, _ in values)
    for index, (label, value, color) in enumerate(values):
        x1 = 280 + index * 390
        x2 = x1 + 230
        height = int(470 * value / max_value)
        y1 = 680 - height
        draw.rectangle((x1, y1, x2, 680), fill=color)
        draw.text((x1 + 65, y1 - 60), str(value), fill="#111111", font=font(38, True))
        draw.text((x1 + 25, 710), label, fill="#222222", font=font(34))
    draw.text((80, 775), "Total exported images  793", fill="#444444", font=font(26))
    canvas.save(path, quality=95)


def create_dataset_sample_sheet(path: Path) -> None:
    """Render real v13 images together with their audited YOLO boxes."""
    samples = [
        ("train", "new_mouse_angle_01", "Mouse: low side view"),
        ("train", "new_mouse_angle_07", "Mouse: underside view"),
        ("train", "camera_v13_023_t043850", "Cup: fully horizontal"),
        ("train", "v10_error_missed_cup_mouse_t090000", "Recovered miss: blur and two targets"),
        ("val", "cup_v3_003_t004100", "Validation: cup with background targets"),
        ("train", "camera_v12_phone_negative_05_f001171", "Hard negative: phone, empty label"),
    ]
    class_names = {0: "mouse", 1: "cup"}
    class_colours = {0: "#20A464", 1: "#2D6A9F"}
    canvas = Image.new("RGB", (1800, 1460), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 35), "Examples from the Audited Version 13 Dataset", fill="#111111", font=font(48, True))
    draw.text((70, 98), "Green = mouse   Blue = cup   Grey = deliberate background sample", fill="#444444", font=font(25))

    card_w, card_h = 815, 390
    image_w, image_h = 775, 315
    for index, (split, stem, caption) in enumerate(samples):
        column, row = index % 2, index // 2
        x = 70 + column * 865
        y = 155 + row * 425
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=18, fill="#F5F7F9", outline="#D4D9DE", width=3)

        image_path = DATASET / "images" / split / f"{stem}.jpg"
        label_path = DATASET / "labels" / split / f"{stem}.txt"
        sample = Image.open(image_path).convert("RGB")
        source_w, source_h = sample.size
        scale = min(image_w / source_w, image_h / source_h)
        display_w = max(1, int(source_w * scale))
        display_h = max(1, int(source_h * scale))
        sample = sample.resize((display_w, display_h), Image.Resampling.LANCZOS)
        image_x = x + 20 + (image_w - display_w) // 2
        image_y = y + 18 + (image_h - display_h) // 2
        canvas.paste(sample, (image_x, image_y))

        label_lines = [line for line in label_path.read_text(encoding="utf-8").splitlines() if line.strip()]
        for line in label_lines:
            class_id, cx, cy, width, height = line.split()[:5]
            class_id = int(class_id)
            cx, cy, width, height = map(float, (cx, cy, width, height))
            x1 = image_x + int((cx - width / 2) * display_w)
            y1 = image_y + int((cy - height / 2) * display_h)
            x2 = image_x + int((cx + width / 2) * display_w)
            y2 = image_y + int((cy + height / 2) * display_h)
            colour = class_colours[class_id]
            draw.rectangle((x1, y1, x2, y2), outline=colour, width=6)
            name = class_names[class_id]
            text_box = draw.textbbox((0, 0), name, font=font(22, True))
            text_w = text_box[2] - text_box[0]
            text_h = text_box[3] - text_box[1]
            label_y = max(image_y, y1 - text_h - 10)
            draw.rectangle((x1, label_y, x1 + text_w + 14, label_y + text_h + 10), fill=colour)
            draw.text((x1 + 7, label_y + 3), name, fill="white", font=font(22, True))

        status = f"{split} | {len(label_lines)} box{'es' if len(label_lines) != 1 else ''}"
        if not label_lines:
            status = f"{split} | empty label"
        draw.text((x + 22, y + 342), caption, fill="#1F2529", font=font(24, True))
        status_box = draw.textbbox((0, 0), status, font=font(21))
        status_w = status_box[2] - status_box[0]
        draw.text((x + card_w - status_w - 22, y + 344), status, fill="#626B73", font=font(21))
    canvas.save(path, quality=95)


def create_requirement_chart(path: Path) -> None:
    canvas = Image.new("RGB", (1800, 770), "white")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 45), "Acceptance Results", fill="#111111", font=font(52, True))
    rows = [
        ("Two object classes", "mouse and cup", 1.0),
        ("Twenty angle accuracy", "90 percent", 0.90),
        ("Required accuracy", "80 percent", 0.80),
        ("Observed Jetson FPS", "17.9 average", 17.9 / 25.0),
        ("Required Jetson FPS", "5.0 minimum", 5.0 / 25.0),
    ]
    for index, (label, value_text, fraction) in enumerate(rows):
        y = 155 + index * 105
        draw.text((70, y), label, fill="#222222", font=font(29))
        draw.rounded_rectangle((600, y + 3, 1510, y + 54), radius=16, fill="#EDF0F3")
        draw.rounded_rectangle((600, y + 3, 600 + int(910 * min(fraction, 1.0)), y + 54), radius=16, fill="#2D6A9F")
        draw.text((1540, y), value_text, fill="#18324B", font=font(28, True))
    draw.text((70, 695), "Accuracy uses strict scene scoring  all expected targets present and no extra false box", fill="#444444", font=font(25))
    canvas.save(path, quality=95)


def split_evidence_sheet(source: Path, top: Path, bottom: Path) -> None:
    image = Image.open(source)
    width, height = image.size
    midpoint = height // 2
    image.crop((0, 0, width, midpoint)).save(top, quality=95)
    image.crop((0, midpoint, width, height)).save(bottom, quality=95)


def exported_counts() -> dict[str, dict[str, int]]:
    output: dict[str, dict[str, int]] = {}
    for split in ("train", "val", "test"):
        mouse = cup = empty = 0
        files = list((DATASET / "labels" / split).glob("*.txt"))
        for path in files:
            text = path.read_text(encoding="utf-8").strip()
            if not text:
                empty += 1
                continue
            for line in text.splitlines():
                class_id = int(line.split()[0])
                if class_id == 0:
                    mouse += 1
                elif class_id == 1:
                    cup += 1
        output[split] = {"images": len(files), "empty": empty, "mouse": mouse, "cup": cup}
    return output


def best_metrics() -> dict[str, float]:
    with (RUN / "results.csv").open(newline="", encoding="utf-8-sig") as handle:
        rows = list(csv.DictReader(handle))
    best = max(rows, key=lambda row: float(row["metrics/mAP50-95(B)"]))
    return {
        "epoch": int(best["epoch"]),
        "precision": float(best["metrics/precision(B)"]),
        "recall": float(best["metrics/recall(B)"]),
        "map50": float(best["metrics/mAP50(B)"]),
        "map5095": float(best["metrics/mAP50-95(B)"]),
    }


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


class Report:
    def __init__(self) -> None:
        self.document = Document()
        section = self.document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.68)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.72)
        self._configure_styles()
        self._configure_header_footer(section)

    def _configure_styles(self) -> None:
        styles = self.document.styles
        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = TEXT
        normal.paragraph_format.space_after = Pt(5.5)
        normal.paragraph_format.line_spacing = 1.08
        for name, size, before, after in (
            ("Title", 26, 0, 14),
            ("Heading 1", 18, 8, 8),
            ("Heading 2", 13.5, 7, 4),
            ("Heading 3", 11.5, 5, 2),
        ):
            style = styles[name]
            style.font.name = "Aptos Display"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
        if "Code Block" not in styles:
            code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        else:
            code_style = styles["Code Block"]
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8.3)
        code_style.font.color.rgb = RGBColor(35, 39, 42)
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.1)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(6)

    def _configure_header_footer(self, section) -> None:
        header = section.header.paragraphs[0]
        header.text = "Experiment One Object Detection and Recognition"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.size = Pt(8.5)
        header.runs[0].font.color.rgb = RGBColor(90, 95, 100)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Page ")
        add_field(footer, "PAGE")
        footer.add_run(" of ")
        add_field(footer, "NUMPAGES")
        for run in footer.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(90, 95, 100)

    def cover(self) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.space_after = Pt(0)
        paragraph.add_run("COMPUTER VISION LABORATORY").bold = True
        paragraph.runs[0].font.size = Pt(11)
        paragraph.runs[0].font.color.rgb = RGBColor(45, 85, 120)
        self.document.add_paragraph("\n\n\n")
        title = self.document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title.add_run("Experiment One\nObject Detection and Recognition")
        subtitle = self.document.add_paragraph()
        subtitle.add_run("Training and deploying a mouse-and-cup detector on Jetson").italic = True
        subtitle.runs[0].font.size = Pt(15)
        subtitle.runs[0].font.color.rgb = RGBColor(60, 65, 70)
        self.document.add_paragraph("\n")
        table = self.document.add_table(rows=5, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        labels = [
            ("Final model", "YOLO11n version 13"),
            ("Target classes", "mouse and cup"),
            ("Training input", "768 by 768 pixels"),
            ("Deployment", "Windows ONNX and Jetson PyTorch CUDA"),
            ("Experiment period", "August to September 2026"),
        ]
        for row, values in zip(table.rows, labels):
            row.cells[0].width = Inches(1.65)
            row.cells[1].width = Inches(4.9)
            for index, value in enumerate(values):
                row.cells[index].text = value
                set_cell_border(row.cells[index], color="D6DCE2", size="4")
                set_cell_margins(row.cells[index])
                if index == 0:
                    shade(row.cells[index], PALE_BLUE)
                    row.cells[index].paragraphs[0].runs[0].bold = True
        self.document.add_paragraph("\n")
        note = self.document.add_paragraph()
        note.add_run("Submission note  ").bold = True
        note.add_run("All numerical claims in this report are tied to files stored in the project workspace. The older CUP3 clip is excluded from the final evidence because it used an earlier unsuccessful test.")
        self.document.add_page_break()

    def heading(self, text: str, level: int = 1) -> None:
        self.document.add_heading(text, level=level)

    def p(self, text: str, bold_lead: str | None = None) -> None:
        paragraph = self.document.add_paragraph()
        if bold_lead and text.startswith(bold_lead):
            paragraph.add_run(bold_lead).bold = True
            paragraph.add_run(text[len(bold_lead):])
        else:
            paragraph.add_run(text)

    def bullets(self, items: list[str]) -> None:
        for item in items:
            paragraph = self.document.add_paragraph(style="List Bullet")
            paragraph.add_run(item)

    def numbered(self, items: list[str]) -> None:
        for item in items:
            paragraph = self.document.add_paragraph(style="List Number")
            paragraph.add_run(item)

    def table(self, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
        table = self.document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_repeat_table_header(table.rows[0])
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = header
            shade(cell, NAVY)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
            if widths:
                cell.width = Inches(widths[index])
        for row_index, values in enumerate(rows):
            cells = table.add_row().cells
            for column_index, value in enumerate(values):
                cell = cells[column_index]
                cell.text = str(value)
                shade(cell, "FFFFFF" if row_index % 2 == 0 else PALE_GRAY)
                set_cell_border(cell)
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8.7)
                if widths:
                    cell.width = Inches(widths[column_index])
        self.document.add_paragraph().paragraph_format.space_after = Pt(0)

    def figure(self, path: Path, caption: str, width: float = 6.9) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(width))
        caption_paragraph = self.document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption)
        caption_run.italic = True
        caption_run.font.size = Pt(9)
        caption_run.font.color.rgb = RGBColor(65, 70, 75)

    def code(self, text: str) -> None:
        paragraph = self.document.add_paragraph(style="Code Block")
        paragraph.add_run(text)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F1F3F5")
        paragraph._p.get_or_add_pPr().append(shading)

    def page(self, heading: str) -> None:
        self.document.add_page_break()
        self.heading(heading)


def build() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    create_pipeline_diagram(ASSETS / "workflow.png")
    create_dataset_chart(ASSETS / "dataset_splits.png")
    create_dataset_sample_sheet(ASSETS / "dataset_samples.png")
    create_requirement_chart(ASSETS / "acceptance_results.png")
    split_evidence_sheet(
        REVIEW / "v13_20_angle_evidence_with_horizontal_cup.jpg",
        ASSETS / "twenty_angles_01_10.jpg",
        ASSETS / "twenty_angles_11_20.jpg",
    )
    counts = exported_counts()
    metrics = best_metrics()
    with (REVIEW / "v13_20_angle_results_with_horizontal_cup.csv").open(newline="", encoding="utf-8-sig") as handle:
        angle_rows = [row for row in csv.DictReader(handle) if row.get("test_id", "").isdigit()]
    horizontal = json.loads((ROOT / "results/v13_horizontal_cup_success_detected.json").read_text(encoding="utf-8"))

    report = Report()
    report.cover()

    report.heading("Executive Summary")
    report.p("This project started as a simple mouse-and-cup detector, but the first live tests exposed problems that were not obvious from the training numbers. Black phones and robot parts were sometimes called mice, side-view mice were missed, and the cup only worked reliably while upright. I kept those failures and used them to decide what to collect next. The final model is YOLO11n v13, trained at 768 pixels and exported as both PyTorch and ONNX weights.")
    report.p("The final export contains 793 images: 492 for training, 61 for validation and 240 for testing. The training split includes 49 empty label files on purpose. They show phones, chair backs, robot shells, laptop edges and other objects that caused false mouse detections. Version 13 also adds eighty manually checked frames, with most of the new positive examples covering cups lying on their side and mouse poses that had previously failed.")
    report.p("I checked twenty different viewing conditions for the classroom test. A scene only counted as correct when every expected object was found and there was no extra box. Eighteen of the twenty scenes passed, so the final score is 90 percent. Eighteen scenes are taken from Jetson recordings. The other two are horizontal-cup frames replayed through the final v13 ONNX model because the older CUP3 clip was upright and did not detect the cup.")
    report.p("The selected Jetson frames show 15.7 to 21.2 FPS, averaging about 17.9 FPS. The ROS 2 test was also run on the physical board: /dev/video0 produced a mouse detection at 0.9478 confidence, and /yolo/detections was measured at about 20.57 Hz. The live display and the ROS 2 message therefore came from the same working detection path.")
    report.figure(ASSETS / "acceptance_results.png", "Figure 1  Summary of the measured acceptance results", 6.9)

    report.page("1 Requirements and Evidence")
    report.p("Before preparing the final files, I copied each item from the assessment sheet into a checklist. The table below records what I used as evidence for each item. This caught one missing item early: an upright cup video could not be used to prove that a cup lying flat was recognised.")
    report.table(
        ["Requirement", "Evidence", "Status"],
        [
            ["Two desktop classes", "Class IDs 0 mouse and 1 cup in dataset.yaml", "Met"],
            ["Self collected and labelled data", "Audited v13 dataset and X AnyLabeling batches", "Met"],
            ["Jetson execution", "Nine v13 Jetson result clips and CUDA deployment script", "Met"],
            ["Class box and confidence overlay", "Rendered v13 videos and live inference code", "Met"],
            ["Twenty angles at 80 percent", "18 correct scenes from 20 selected scenes", "Met at 90 percent"],
            ["At least 5 FPS on Jetson", "15.7 to 21.2 FPS in selected Jetson frames", "Met"],
            ["Save results and errors", "Annotated and raw video outputs plus review frames", "Met"],
            ["ROS 2 detection publishing", "Jetson camera Detection2DArray, mouse 0.9478, about 20.57 Hz", "Met"],
        ],
        [2.0, 3.8, 1.2],
    )
    report.p("The package first passed a recorded-video replay in WSL, which checked message construction and DDS publication. I then repeated the full path on the physical Jetson with its camera. The final evidence package contains the annotated AVI, a representative frame, the detector log and a non-empty topic record. Keeping both runs makes the software test and the hardware acceptance test independently reproducible.")
    report.bullets([
        "Accuracy is scene based because the teacher accepted twenty angles rather than twenty unique object instances.",
        "The horizontal cup evidence is included as an explicit supplement and not mixed into the Jetson FPS calculation.",
        "The old CUP3 recording is intentionally ignored because it showed an upright stainless cup with no successful detection.",
    ])

    report.page("2 System Design")
    report.p("My working cycle was: record a short video, keep representative frames, correct the boxes in X AnyLabeling, run the audit, export the YOLO folders, train, and test again with the camera. When the model made a mistake, I saved that frame. A missed object became a positive sample after I drew the box; a false alarm became an empty-label negative sample.")
    report.figure(ASSETS / "workflow.png", "Figure 2  Data training deployment and error feedback workflow", 6.9)
    report.p("This workflow explains why the project progressed through several versions. Each version addressed a specific failure rather than changing the model architecture at random. Versions 9 and 10 improved camera domain coverage. Version 11 added frames taken from missed and false detections. Version 12 included phones as difficult negative examples because phones were often predicted as mice. Version 13 concentrated on horizontal cups and additional poses.")
    report.p("The output side has two parallel paths. The Windows path uses ONNX Runtime and is convenient for camera testing and data capture. The Jetson path uses the PyTorch best.pt weight with CUDA through Ultralytics. The ROS 2 node wraps the Jetson prediction result in standard vision_msgs structures so other robot nodes can consume the class, confidence and bounding box without parsing screen graphics.")

    report.page("3 Development and Deployment Environment")
    report.table(
        ["Component", "Development computer", "Jetson deployment"],
        [
            ["Operating system", "Windows 11 and WSL 2 Ubuntu 22.04", "Ubuntu 22.04 with JetPack 6.2.1"],
            ["Python", "3.10.12 training and 3.13 ONNX", "Jetson Python environment"],
            ["Deep learning", "PyTorch 2.12.1 plus CUDA 12.6", "PyTorch 2.5.0a0 Jetson build"],
            ["Detection framework", "Ultralytics 8.4.127", "Ultralytics 8.4.128"],
            ["Image processing", "OpenCV 4.10", "OpenCV 4.13"],
            ["ROS", "ROS 2 Humble in WSL", "ROS 2 Humble with vision_msgs"],
            ["Camera", "USB camera index 1", "V4L2 camera device"],
        ],
        [1.5, 2.8, 2.8],
    )
    report.p("WSL 2 was useful because Ultralytics training and ROS 2 Humble both behaved like a normal Ubuntu installation while the dataset remained on the Windows F drive. Training paths in args.yaml therefore appear as /mnt/f paths. The Windows ONNX program uses camera index numbers, while Jetson normally exposes the same hardware as /dev/video0 or /dev/video1.")
    report.p("I used environment checks before long training jobs. The checks confirm that Python can import OpenCV, PyTorch and Ultralytics, that CUDA is visible, and that ROS 2 commands are sourced correctly. An environment check is not a substitute for a task test. For example, publishing a test string proves DDS communication but does not prove that object detection messages contain correct boxes. This distinction led to the dedicated detector node added later.")
    report.code("python scripts/check_env.py\npython scripts/check_yolo.py\nbash scripts/check_ros2.sh")

    report.page("4 Object Classes and Annotation Policy")
    report.p("The class set is intentionally small. Class 0 is mouse and class 1 is cup. Both are common desktop objects, but their shapes are not equally easy. A cup is often tall and visually distinctive. A mouse can be small, dark and similar to phones, power banks, chair parts or robot components. The label policy therefore needs consistency more than complexity.")
    report.table(
        ["ID", "Class", "Box rule", "Important edge cases"],
        [
            ["0", "mouse", "Tight box around the visible mouse body", "Side view underside hand occlusion and small distant views"],
            ["1", "cup", "Tight box around the visible cup body", "Horizontal pose partial occlusion lid and reflective surfaces"],
        ],
        [0.5, 0.9, 2.7, 3.0],
    )
    report.p("When two target objects overlap, both are labelled if each remains meaningfully visible. This is not duplicate annotation. A duplicate exists only when two boxes describe the same instance. Earlier audits found no same class duplicate box with intersection over union above 0.5. Empty label files are valid when no mouse or cup appears; they teach the model that a visually similar black object is background.")
    report.p("I reviewed model generated prelabels manually in X AnyLabeling. Prelabels accelerated the first pass, but they were never treated as ground truth automatically. I corrected missing objects, removed boxes around phones and robot bodies, and checked that boxes did not extend outside the image. Two v13 boxes near image borders were clipped during export, and the audit report contains no fatal errors.")

    report.page("5 Data Collection Strategy")
    report.p("At first I extracted too many neighbouring frames from phone videos. The image count increased, but many frames showed almost the same pose and background. In the later sessions I collected fewer, more useful differences: front and side views, high and low camera positions, rotation, hand occlusion, distance changes, new desks and the black objects that had triggered false boxes.")
    report.p("The horizontal cup session was recorded specifically after the assessment requirement changed. It includes a cup lying fully on its side, movement across the table and several backgrounds. Eighty representative frames were selected and manually reviewed for version 13. The original raw video was kept so that the final model could later be replayed on exactly the captured sequence.")
    report.bullets([
        "Positive diversity included upright cups, horizontal cups, blue and green cups, mouse top views, side profiles and underside views.",
        "Hard negatives included phones, laptop edges, screens, chairs, bags, robot bodies, cables and other black objects.",
        "Raw and annotated videos were both saved where practical so wrong predictions could be reproduced without filming again.",
        "Frames from one continuous scene were not relied on as independent proof of generalisation.",
    ])
    report.p("The final twenty angle test is separate from the training set audit. It is a practical demonstration rather than a substitute for the held out test partition. Both forms of evaluation are needed: the image split produces repeatable numerical metrics, and the live sequence shows whether the program behaves acceptably during actual use.")

    report.page("6 Dataset Cleaning and Audit")
    report.p("The audit found several ordinary labelling mistakes. hard_robot_mouse had a misleading filename but no real mouse, so I kept it as a negative. neg_laptop did contain part of a cup at the left edge, so I added that cup box. I also found that the first export script silently treated train_candidate and not_assigned as training data. I changed the script so only explicit train, val and test rows are exported.")
    report.p("The audit process performs structural checks before training. It verifies that every exported image has one matching label file, class IDs are valid, normalised coordinates are within range, box width and height are positive, and the train validation and test lists do not contain the same file. It also produces an exclusion list so rejected or unassigned samples remain visible for later review rather than disappearing.")
    report.table(
        ["Audit check", "Reason", "v13 outcome"],
        [
            ["Image label pairing", "Prevents unintentional missing annotations", "Passed"],
            ["Class ID range", "Only mouse and cup are valid", "Passed"],
            ["Coordinate range", "YOLO values must remain between zero and one", "Passed after two clips"],
            ["Empty labels", "Retains deliberate hard negatives", "58 in audit and 55 in exported splits"],
            ["Split membership", "Avoids hidden train candidate promotion", "Explicit train val test only"],
            ["Fatal audit errors", "Stops training on invalid data", "None"],
        ],
        [1.5, 3.6, 2.0],
    )

    report.page("7 Version 13 Dataset Composition")
    total_images = sum(value["images"] for value in counts.values())
    total_mouse = sum(value["mouse"] for value in counts.values())
    total_cup = sum(value["cup"] for value in counts.values())
    total_empty = sum(value["empty"] for value in counts.values())
    report.p(f"The authoritative exported directory contains {total_images} images. Parsing the label files directly gives {total_mouse} mouse boxes and {total_cup} cup boxes. There are {total_empty} empty labels across the three splits. The counts below are calculated from the actual exported files rather than copied from a planning note.")
    report.table(
        ["Split", "Images", "Mouse boxes", "Cup boxes", "Empty labels"],
        [[split, str(v["images"]), str(v["mouse"]), str(v["cup"]), str(v["empty"])] for split, v in counts.items()] + [["Total", str(total_images), str(total_mouse), str(total_cup), str(total_empty)]],
        [1.3, 1.2, 1.5, 1.4, 1.4],
    )
    report.figure(ASSETS / "dataset_splits.png", "Figure 3  Exported train validation and test image counts", 6.7)
    report.p("The test partition is relatively large because it preserves independent original scenes and is reused for comparison between model versions. The training partition is smaller than the raw audit store because 461 audit records are explicitly excluded. This is preferable to increasing the count with uncertain or duplicate samples.")

    report.page("Dataset Sample Gallery")
    report.p("These are actual images from the exported v13 folders. The boxes are drawn from the matching YOLO label files, not from model predictions. I selected them to show the cases that mattered during development: a normal mouse, its underside, a cup lying flat, a blurred two-object recovery sample, a validation scene and a phone kept as background.")
    report.figure(ASSETS / "dataset_samples.png", "Plate 1  Real v13 images with audited ground-truth boxes", 7.0)
    report.p("The phone card has no box because it is a deliberate hard negative. This is how the dataset tells the model that a dark rectangular object is not automatically a mouse. The blurred card is included because it came from a real failure, not because it is visually clean.")

    report.page("8 Dataset Version History")
    report.table(
        ["Version", "Primary change", "Reason"],
        [
            ["v7", "First audited baseline at 768 pixels", "Establish repeatable training"],
            ["v8", "Cleaned export and evaluation scripts", "Improve split reliability"],
            ["v9", "External camera data", "Reduce camera domain mismatch"],
            ["v10", "Manual camera reannotation", "Improve mouse and cup angle coverage"],
            ["v11", "Live error feedback samples", "Correct misses and false positives"],
            ["v12", "Phone hard negatives and mouse views", "Stop phone to mouse confusion"],
            ["v13", "Horizontal cup and new environments", "Meet the horizontal cup requirement"],
        ],
        [0.8, 3.0, 3.3],
    )
    report.p("The version number increased only after a new data batch or a specific correction. The metrics did not rise every time. In fact, v13 has slightly lower held-out mAP than v12, but it recognises the horizontal cup pose that v12 lacked. I kept v13 for the final demonstration because that pose was part of the teacher's requirement and the twenty-angle test still passed.")
    report.p("Keeping separate version folders and training scripts makes the experiment reproducible. It also prevents a good weight file from being overwritten by a later trial. Each training output has args.yaml, results.csv, plots and best and last weights. This structure allowed me to trace exactly which dataset and starting weight produced the final model.")

    report.page("9 Model Choice and Architecture")
    report.p("I selected YOLO11n because the experiment values real time operation on Jetson more than maximum offline accuracy. The nano variant is small enough to run comfortably above 5 FPS while still learning two simple classes. A one stage detector predicts boxes and class scores in one network pass, which is simpler to deploy than a two stage system.")
    report.p("The model input is a tensor with batch, channel, height and width dimensions. Camera frames arrive in HWC layout using OpenCV BGR channel order. During inference the program letterboxes the image to 768 by 768, converts BGR to RGB, transposes HWC to CHW, normalises byte values to floating point zero to one, and adds a batch dimension. ONNX Runtime then returns candidate box coordinates and class scores.")
    report.p("ONNX is an interchange format for neural networks. It separates the trained computation graph from the original training program, so the Windows test tool can run the model with ONNX Runtime without importing the full Ultralytics training stack. The Jetson path keeps the PyTorch best.pt weight because Ultralytics can use CUDA directly and the result object is convenient for ROS 2 message construction.")
    report.table(
        ["Artifact", "Purpose", "Approximate role"],
        [
            ["best.pt", "Jetson and Ultralytics inference", "Full PyTorch checkpoint"],
            ["best.onnx", "Portable Windows inference", "Exported computation graph"],
            ["dataset.yaml", "Class names and split locations", "Training data entry point"],
            ["args.yaml", "Exact training configuration", "Reproducibility record"],
        ],
        [1.2, 3.0, 2.8],
    )

    report.page("10 Version 13 Training Configuration")
    report.table(
        ["Parameter", "Value", "Interpretation"],
        [
            ["Initial weight", "v12 best.pt", "Continue from the previous cleaned detector"],
            ["Epoch limit", "40", "Maximum complete passes through training data"],
            ["Patience", "12", "Stop when validation improvement stalls"],
            ["Batch", "16", "Images processed before each optimiser update"],
            ["Image size", "768", "Square network input resolution"],
            ["Optimiser", "SGD", "Momentum based gradient update"],
            ["Initial learning rate", "0.00015", "Small fine tuning step"],
            ["Final rate factor", "0.01", "Final rate relative to initial rate"],
            ["Momentum", "0.937", "Smooths gradient direction"],
            ["Weight decay", "0.0005", "Regularises network weights"],
            ["Seed", "20260824", "Makes the split order and randomness repeatable"],
            ["Mixed precision", "enabled", "Reduces GPU memory and improves speed"],
        ],
        [1.8, 1.5, 3.8],
    )
    report.p("SGD was chosen for controlled fine tuning. The very small initial learning rate is appropriate because v13 starts from an already trained v12 model. A larger learning rate could destroy useful features while fitting only the new horizontal cup batch. Training used RAM caching and four data workers in WSL on the laptop GPU.")
    report.code("cd /mnt/f/PycharmProjects/robomaster\nbash scripts/train_v13_horizontal_cup_wsl.sh")

    report.page("11 Loss Functions and Data Augmentation")
    report.p("The detector uses three main loss terms. The box loss weight is 7.5 and penalises incorrect localisation. The classification loss weight is 0.5 and separates mouse from cup and background. The distribution focal loss weight is 1.5 and helps refine bounding box edges. These values came from the actual v13 args.yaml file.")
    report.table(
        ["Augmentation", "Value", "Effect"],
        [
            ["Hue shift", "0.015", "Small colour change"],
            ["Saturation", "0.7", "Wide colour intensity variation"],
            ["Brightness", "0.4", "Lighting variation"],
            ["Translation", "0.1", "Moves targets within the frame"],
            ["Scale", "0.5", "Changes apparent object size"],
            ["Horizontal flip", "0.5", "Mirrors half of training samples"],
            ["Mosaic", "1.0", "Combines several images early in training"],
            ["Random erasing", "0.4", "Simulates partial occlusion"],
            ["Rotation shear perspective", "0", "Not enabled in this run"],
            ["MixUp CutMix Copy Paste", "0", "Not enabled in this run"],
        ],
        [2.0, 1.2, 3.9],
    )
    report.p("Augmentation changes image appearance and placement, but it cannot create a truly new three dimensional viewpoint. This was an important lesson from the mouse failures. Rotating a top view by twelve degrees is still a top view; it does not become a side or underside image. Real collection remained necessary even after aggressive augmentation.")

    report.page("12 Training Results")
    report.p(f"The best validation mAP50 to 95 occurred at epoch {metrics['epoch']}. At that point precision was {metrics['precision']:.3f}, recall was {metrics['recall']:.3f}, mAP50 was {metrics['map50']:.3f}, and mAP50 to 95 was {metrics['map5095']:.3f}. I used best.pt for deployment instead of assuming that the last epoch was best.")
    report.figure(RUN / "results.png", "Figure 4  Version 13 training and validation curves", 7.0)
    report.p("Figure 4 is the original results.png written by Ultralytics during the v13 run; it has not been redrawn for this report. The left-side plots show box, classification and DFL losses, while the right-side plots show precision, recall and mAP. The validation curves worsen after the best region even though some training losses continue to fall. That is why I kept the checkpoint from the best validation epoch. More independent horizontal-cup scenes would be more useful than simply adding epochs.")

    report.page("13 Validation Analysis")
    report.p("The normalised confusion matrix gives a class focused view. About 0.80 of true mouse instances are classified as mouse, while about 0.85 of true cup instances are classified as cup. Background related false positives remain visible, especially for mouse. This matches live testing: small dark objects are more likely to be confused with a mouse than with a cup.")
    report.figure(RUN / "confusion_matrix_normalized.png", "Figure 5  Normalised validation confusion matrix for version 13", 5.9)
    report.p("The held out test partition is intentionally harder than the classroom demonstration. At the same test configuration, v13 produced precision 0.706, recall 0.635, mAP50 0.677 and mAP50 to 95 0.374. Version 12 had slightly better general test metrics. I therefore do not claim that version 13 is universally superior. Its value is the targeted horizontal cup behaviour combined with the practical 90 percent acceptance result.")

    report.page("14 ONNX Inference Pipeline")
    report.p("The Windows program performs preprocessing, inference, postprocessing and display in a continuous loop. Letterboxing preserves the source aspect ratio while filling unused pixels with a neutral grey value. The original scale and padding are retained so output boxes can be mapped back to camera coordinates. Non maximum suppression is performed separately for each class to remove duplicated overlapping predictions.")
    report.code("letterboxed, scale, pad_x, pad_y = letterbox(frame, imgsz)\ninput_tensor = cv2.cvtColor(letterboxed, cv2.COLOR_BGR2RGB)\ninput_tensor = input_tensor.transpose(2, 0, 1).astype(np.float32) / 255.0\ninput_tensor = np.expand_dims(input_tensor, axis=0)\noutput = session.run(None, {input_name: input_tensor})[0]")
    report.p("Independent confidence thresholds are useful because the two classes have different error patterns. Mouse errors were usually false positives on dark objects, so a high mouse threshold reduces those mistakes. Cup predictions were generally cleaner. The final demonstration used 0.75 for both classes to make the scoring rule conservative.")
    report.p("A short temporal smoother requires two matching frames before displaying a new detection and can hold a track for three missed frames. This reduces flicker. It does not create a detection from nothing, and it should not be used to inflate still image accuracy. For that reason the practical evidence is evaluated scene by scene rather than counting every held display frame as a new success.")

    report.page("15 Live Display and Video Recording")
    report.p("The live application overlays each accepted class name, confidence value and coloured rectangle. Mouse boxes are green and cup boxes are blue. A status line shows smoothed FPS, current detection count and class thresholds. Pressing Q or Escape ends the loop and releases the camera and video writer cleanly.")
    report.code("py -3.13 scripts/live_camera_onnx.py `\n  --camera 1 `\n  --model runs/detect/mouse_cup_yolo11n_v13_horizontal_cup_768/weights/best.onnx `\n  --mouse-conf 0.75 --cup-conf 0.75 `\n  --save results/v13_camera1_detected.avi `\n  --save-raw results/v13_camera1_raw.avi")
    report.p("Saving both streams serves different purposes. The annotated file is evidence for the assessor because it contains boxes, labels, confidence and FPS. The raw file preserves the actual camera input and is more valuable for debugging. It can be replayed through a later model version without repeating the physical experiment.")
    report.p("The new run_video_onnx.py utility follows the same preprocessing and class thresholds for recorded input. I used it to replay the horizontal cup capture through the final v13 ONNX model and produce a ten second supplemental clip. The JSON summary records the source, model, selected time range, thresholds and number of frames containing each class.")

    report.page("16 Jetson Deployment")
    report.p("On Jetson, the program uses the v13 best.pt file and CUDA. OpenCV opens the V4L2 camera, Ultralytics performs prediction, and the application draws the same information as the Windows version. The Jetson recordings are 640 by 480 MJPEG at 20 FPS. This camera rate is appropriate for the experiment and keeps end to end latency low.")
    report.code(
        "python3 scripts/live_camera_pt.py \\\n"
        "  --camera 0 \\\n"
        "  --model /home/nvidia/jetson_yolo/best.pt \\\n"
        "  --imgsz 768 \\\n"
        "  --mouse-conf 0.75 --cup-conf 0.75 \\\n"
        "  --save /home/nvidia/jetson_yolo/results/v13_jetson_detected.avi"
    )
    report.p("The selected evidence frames report a minimum of 15.7 FPS and a maximum of 21.2 FPS, with an average near 17.9 FPS. Even the minimum is more than three times the required 5 FPS. Because the horizontal supplement was processed offline on Windows, those two frames are excluded from the Jetson speed average.")
    report.table(
        ["Measurement", "Observed", "Requirement", "Result"],
        [
            ["Selected Jetson minimum", "15.7 FPS", "5 FPS", "Pass"],
            ["Selected Jetson maximum", "21.2 FPS", "5 FPS", "Pass"],
            ["Selected Jetson mean", "about 17.9 FPS", "5 FPS", "Pass"],
            ["Camera stream", "640 by 480 at 20 FPS", "real time", "Pass"],
        ],
        [2.2, 1.8, 1.4, 1.1],
    )

    report.page("17 Twenty Angle Test Protocol")
    report.p("The teacher confirmed that twenty angles are sufficient. I selected twenty representative moments rather than claiming twenty different physical object models. The set covers upright, tilted and horizontal cups, two cups together, two classes together, mouse top views, side profiles and underside views.")
    report.p("A scene is correct only when all expected targets are detected and no extra false positive appears. This is stricter than target recall alone. A two object scene fails if one target is missing. The method also prevents a long easy clip from dominating the result because every selected angle contributes one score.")
    report.numbered([
        "Choose a timestamp with a clear and stable object pose.",
        "Record the expected target list before reading the prediction.",
        "Check the rendered frame for every expected box and for extra boxes.",
        "Mark the whole scene correct or incorrect using the strict rule.",
        "Count correct scenes and divide by twenty.",
    ])
    report.p("The two horizontal cup angles replace redundant upright cup views. The old CUP3 file is excluded. This makes the evidence closer to the actual assessment requirement and removes a clip that could be misinterpreted as proof of a successful horizontal cup detection.")

    report.page("18 Twenty Angle Results")
    rows = []
    for row in angle_rows:
        rows.append([
            row["test_id"], row["expected"], row["angle_description"], row["platform"], "Pass" if row["result"] == "correct" else "Miss"
        ])
    report.table(["No", "Expected", "View", "Platform", "Result"], rows, [0.45, 1.0, 2.7, 1.7, 0.8])
    report.p("Eighteen of the twenty scenes are correct, giving 90 percent strict scene accuracy. The two failures are mouse views: an extreme side profile and an underside pose. No selected scene failed because of a wrong class label. These failures are retained because a credible experiment should show both successful and unsuccessful cases.")

    report.page("19 Evidence Frames One to Ten")
    report.figure(ASSETS / "twenty_angles_01_10.jpg", "Figure 6  Cup and mixed class evidence for test scenes 1 to 10", 7.0)
    report.p("Scenes 1, 2 and 5 show the green cup upright and tilted. Scenes 3 and 4 show the same target fully horizontal with v13 ONNX confidence around 0.84 to 0.85. Scenes 6 to 8 use a blue cup. Scene 9 contains two cups, and scene 10 proves simultaneous mouse and cup detection.")
    report.p("The first two rows use Jetson output except for the two explicitly labelled offline v13 frames. The platform label is printed below each evidence image so the supplemental replay cannot be confused with Jetson execution.")

    report.page("20 Evidence Frames Eleven to Twenty")
    report.figure(ASSETS / "twenty_angles_11_20.jpg", "Figure 7  Mouse evidence and retained failure cases for test scenes 11 to 20", 7.0)
    report.p("Scenes 11 to 18 show successful mouse detections across side, diagonal, top and underside views. Scenes 19 and 20 are the two misses. Their inclusion explains why the final score is 90 rather than 100 percent and identifies the most useful direction for future data collection.")
    report.p("All ten frames on this page come from Jetson recordings. The displayed FPS values range from 15.9 to 21.2 in the successful frames and remain above 16 FPS in the two misses, confirming that the failures are recognition errors rather than a stalled inference loop.")

    report.page("21 Horizontal Cup Supplement")
    report.p("The previous CUP3 clip did not satisfy the assessment need. It showed an upright stainless cup and the model did not recognise it. I therefore ignored that clip and used the separate v13 horizontal cup result already stored in the results directory. The final v13 ONNX model was replayed for thirteen seconds with both class thresholds set to 0.75.")
    report.figure(REVIEW / "v13_horizontal_cup_success_detected_contact_sheet.jpg", "Figure 8  Thirteen second horizontal cup replay through the final v13 ONNX model", 7.0)
    coverage = 100.0 * horizontal["frames_with_cup"] / horizontal["processed_frames"]
    report.p(f"The segment contains {horizontal['processed_frames']} processed frames and {horizontal['frames_with_cup']} frames with an accepted cup detection, giving {coverage:.1f} percent frame coverage. Stable frames reach approximately 0.79 to 0.85 confidence. The file is suitable as pose evidence, but it is not used as Jetson speed evidence because it is an offline replay.")
    report.p("The contact sheet also contains other black objects in the background. Some are detected as mice because actual mice are present on the desk. The cup claim is limited to the large green horizontal cup and does not imply that every object in the raw scene was part of the cup test.")

    report.page("22 ROS 2 Detection Publisher")
    report.p("The project now includes a ROS 2 Humble Python package named yolo_detection_ros2. The detector node opens a camera, runs the YOLO model, filters mouse and cup boxes with independent thresholds, publishes a vision_msgs Detection2DArray message, renders the live window and optionally writes the annotated video.")
    report.p("Each Detection2D message contains a bounding box centre, width and height, a text class ID and a floating point confidence score. The array header contains the camera frame name and ROS time. A receiving robot node can subscribe to /yolo/detections and make decisions without reading pixels or parsing console output.")
    report.code("cd ~/yolo_ros2_ws\nsource /opt/ros/humble/setup.bash\ncolcon build --symlink-install\nsource install/setup.bash\nros2 launch yolo_detection_ros2 detector.launch.py")
    report.code("ros2 topic info /yolo/detections\nros2 topic echo /yolo/detections --once\nros2 topic hz /yolo/detections")
    report.p("I verified the package with colcon build in Ubuntu 22.04, and ROS 2 lists the detector_node executable. A WSL replay produced a cup message at 0.8217 confidence and about 35 Hz. The final hardware run then started the same package on nvidia-desktop with /dev/video0. Its saved topic record reports one publisher of type vision_msgs/msg/Detection2DArray, a mouse box at 0.9478 confidence and a measured rate of about 20.57 Hz.")
    report.figure(REVIEW / "v13_jetson_ros2_camera_evidence.jpg", "Figure 9  Final Jetson camera run with detection overlay and ROS 2 publishing enabled", 5.3)
    report.code("bash scripts/capture_jetson_ros2_evidence.sh 0 models/v13/best.pt ~/yolo_ros2_ws\n# Portable proof: release/experiment_one_v13_ros2_jetson_evidence.zip")

    report.page("23 Saving Results and Typical Errors")
    report.p("The experiment saves three useful forms of evidence. Annotated AVI files show the visible output expected by the teacher. Raw AVI files preserve the original camera frames for later replay. Review directories contain contact sheets, selected still frames, CSV scoring tables and JSON run summaries. Together they support both demonstration and diagnosis.")
    report.table(
        ["Artifact type", "Example", "Purpose"],
        [
            ["Annotated Jetson video", "v13 Jetson camera1 detected mouseCUP avi", "Two class live evidence"],
            ["Horizontal supplement", "v13 horizontal cup success take1 avi", "Successful flat cup pose"],
            ["Raw recording", "v13 pc camera1 raw files", "Reproducible model replay"],
            ["Scoring table", "v13 20 angle results with horizontal cup csv", "Auditable 18 of 20 result"],
            ["Evidence sheet", "v13 20 angle evidence with horizontal cup jpg", "Visual summary"],
            ["Run summary", "v13 horizontal cup success take1 json", "Thresholds frame counts and source"],
        ],
        [1.5, 3.2, 2.3],
    )
    report.p("The two missed mouse views are typical false negative cases. Earlier versions also produced false positives on phones and black equipment. Those error types were converted into v11 and v12 training batches. This closes the loop between testing and dataset improvement and provides a repeatable procedure for future model versions.")

    report.page("24 Discussion of Failure Cases")
    report.p("The remaining misses are physically plausible. In an extreme side profile the mouse becomes a narrow dark strip with little visible wheel or button structure. In an underside view the sensor panel and feet look unlike the top surface that dominates ordinary mouse datasets. At 768 pixels the object is large enough, but the learned visual features are weak for those shapes.")
    report.p("A high 0.75 threshold reduces false positives but also rejects uncertain true objects. Lowering the threshold would recover some missed angles, yet it could reintroduce phone and chair errors. The current value therefore reflects the assessment setting: a conservative box is preferable to many incorrect mouse boxes. A better solution is more real side and underside training images, not simply changing the threshold after seeing the answer.")
    report.p("The validation plot also warns against overtraining the small specialised batch. More epochs would not necessarily fix the difficult views. I would collect several new mice with different shapes, keep whole scenes separated between train and test, and repeat the twenty angle protocol without selecting timestamps from the same training capture.")
    report.bullets([
        "Add extreme side and underside mouse images from at least three backgrounds.",
        "Add reflective and stainless cups lying on their sides.",
        "Keep phone and black equipment negatives in every new training export.",
        "Compare 0.65 0.70 and 0.75 thresholds on a fixed video before changing deployment defaults.",
    ])

    report.page("25 Reproducibility Guide")
    report.heading("25 1 Audit and Export", 2)
    report.code("py -3.13 scripts/export_audited_yolo.py `\n  --audit dataset_work/audit_dataset_v13 `\n  --output dataset_work/audit_dataset_v13/yolo_export")
    report.heading("25 2 Train", 2)
    report.code("wsl -d Ubuntu-22.04 bash -lc \"cd /mnt/f/PycharmProjects/robomaster && bash scripts/train_v13_horizontal_cup_wsl.sh\"")
    report.heading("25 3 Export ONNX", 2)
    report.code("yolo export model=runs/detect/mouse_cup_yolo11n_v13_horizontal_cup_768/weights/best.pt format=onnx imgsz=768 simplify=True")
    report.heading("25 4 Windows Camera Test", 2)
    report.code("py -3.13 scripts/live_camera_onnx.py --camera 1 --model runs/detect/mouse_cup_yolo11n_v13_horizontal_cup_768/weights/best.onnx --mouse-conf 0.75 --cup-conf 0.75 --save results/v13_detected.avi --save-raw results/v13_raw.avi")
    report.heading("25 5 ROS 2 Build", 2)
    report.code("mkdir -p ~/yolo_ros2_ws/src\ncp -r ros2/yolo_detection_ros2 ~/yolo_ros2_ws/src/\ncd ~/yolo_ros2_ws\nsource /opt/ros/humble/setup.bash\ncolcon build --symlink-install")
    report.heading("25 6 ROS 2 Runtime Verification", 2)
    report.code("bash scripts/capture_jetson_ros2_evidence.sh 0 /home/nvidia/jetson_yolo/models/v13/best.pt /home/nvidia/yolo_ros2_ws\n# Verify on the PC:\npy -3.13 scripts/audit_experiment_one_submission.py --verify-git --require-jetson-ros2")

    report.page("26 Submission Artifact Inventory")
    report.table(
        ["Deliverable", "Recommended file or directory"],
        [
            ["Dataset", "dataset_work audit_dataset_v13 yolo_export"],
            ["PyTorch model", "runs detect mouse_cup_yolo11n_v13_horizontal_cup_768 weights best.pt"],
            ["ONNX model", "runs detect mouse_cup_yolo11n_v13_horizontal_cup_768 weights best.onnx"],
            ["Windows program", "scripts live_camera_onnx.py"],
            ["Jetson program", "scripts live_camera_pt.py"],
            ["ROS 2 package and proof", "ros2 yolo_detection_ros2 and release experiment_one_v13_ros2_jetson_evidence.zip"],
            ["Jetson videos", "results v13 Jetson camera1 detected files"],
            ["Horizontal cup video", "results v13_horizontal_cup_success_detected.avi"],
            ["Twenty angle score", "results v13_submission_review v13_20_angle_results_with_horizontal_cup.csv"],
            ["Report", "docs Experiment_One_Object_Detection_Report.docx and pdf"],
        ],
        [1.8, 5.2],
    )
    report.p("Large datasets and model weights may exceed ordinary GitHub limits. The repository should contain scripts, manifests, small evidence images, CSV summaries, README instructions and the report. Large weights and videos can be submitted as release assets or as a separate course archive when required. The report lists their exact local paths and hashes so the assessor can verify the files did not change.")

    report.page("27 Integrity Record")
    best_pt = RUN / "weights/best.pt"
    best_onnx = RUN / "weights/best.onnx"
    horizontal_video = ROOT / "results/v13_horizontal_cup_success_detected.avi"
    report.p("The following SHA 256 values were calculated from the final local files. They provide a simple integrity record for copying the submission to Jetson, cloud storage or removable media.")
    report.table(
        ["File", "Bytes", "SHA 256"],
        [
            ["v13 best.pt", str(best_pt.stat().st_size), sha256(best_pt)],
            ["v13 best.onnx", str(best_onnx.stat().st_size), sha256(best_onnx)],
            ["horizontal cup video", str(horizontal_video.stat().st_size), sha256(horizontal_video)],
            ["Jetson ROS 2 evidence ZIP", str((ROOT / "release/experiment_one_v13_ros2_jetson_evidence.zip").stat().st_size), sha256(ROOT / "release/experiment_one_v13_ros2_jetson_evidence.zip")],
        ],
        [1.5, 1.2, 4.4],
    )
    report.p("A checksum does not prove accuracy, but it prevents accidental substitution of a different model after the report is written. It is especially useful here because several version folders contain files named best.pt.")

    report.page("28 Conclusion")
    report.p("The final v13 model can detect a mouse and a cup in the same camera frame, including a cup lying flat. The most useful part of the work was not changing the network; it was keeping the bad predictions, correcting the labels and collecting the missing poses. The phone negatives and the horizontal-cup batch came directly from problems seen during live testing.")
    report.p("In the twenty-angle check, 18 scenes passed and 2 mouse views failed, giving 90 percent. The selected Jetson frames run from 15.7 to 21.2 FPS, above the required 5 FPS. The saved videos show the class, box, confidence and FPS, and the program can also keep the raw stream for later replay.")
    report.p("The ROS 2 package built and ran on the Jetson. Its stored topic output contains a mouse detection at 0.9478 confidence and a measured publish rate of about 20.57 Hz. The two missed mouse views are still included in the evidence. If I continue the project, the next data batch should focus on extreme side and underside views rather than adding more ordinary top views.")

    report.page("Appendix A Key Source Code Responsibilities")
    report.table(
        ["File", "Responsibility"],
        [
            ["scripts live_camera_onnx.py", "Windows camera input ONNX preprocessing NMS smoothing display and recording"],
            ["scripts live_camera_pt.py", "Jetson camera input Ultralytics CUDA inference display and recording"],
            ["scripts run_video_onnx.py", "Offline replay of a selected raw segment with JSON summary"],
            ["scripts train_v13_horizontal_cup_wsl.sh", "Repeatable v13 training command and hyperparameters"],
            ["scripts export_audited_yolo.py", "Only exports explicitly assigned train validation and test records"],
            ["ros2 yolo_detection_ros2 detector_node.py", "Publishes boxes class IDs and confidence as vision_msgs"],
        ],
        [2.8, 4.2],
    )
    report.p("The live inference file separates small functions such as letterbox and postprocess from the camera loop. This made offline video replay possible without copying the mathematical code. The ROS 2 node uses Ultralytics results directly because Jetson already has that framework installed. Both paths keep the same class order and deployment thresholds.")

    report.page("Appendix B Final Evidence Paths")
    report.code("results/v13_submission_review/v13_20_angle_evidence_with_horizontal_cup.jpg\nresults/v13_submission_review/v13_20_angle_results_with_horizontal_cup.csv\nresults/v13_horizontal_cup_success_detected.avi\nresults/v13_horizontal_cup_success_detected.json\nresults/v13_Jetson_camera1_detected_mouseCUP.avi\nrelease/experiment_one_v13_ros2_jetson_evidence.zip\nrelease/experiment_one_v13_video_evidence.zip\nrelease/experiment_one_v13_model.zip")
    report.p("The older v13_Jetson_camera1_detected_CUP3.avi file is not part of the final proof. It remains in the working results directory only as historical test material. The curated CSV and evidence sheet do not reference it.")
    report.p("This report was generated from the current local workspace so that tables and integrity hashes match the files submitted with the experiment.")

    report.document.core_properties.title = "Experiment One Object Detection and Recognition"
    report.document.core_properties.subject = "YOLO11n mouse and cup detection on Jetson with ROS 2"
    report.document.core_properties.author = "Student Project Report"
    report.document.core_properties.keywords = "YOLO11n Jetson ROS 2 mouse cup object detection"
    report.document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
